"""
DOCX Resume Generator using Template
Replaces placeholders in the template with tailored content
"""

import os
from typing import Dict, List, Any
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess
import tempfile
from datetime import datetime

class DocxResumeGenerator:
    def __init__(self, template_path: str = "data/placeholder_resume.docx"):
        """Initialize with template path"""
        self.template_path = template_path
        
    def generate_resume(self, data: Dict[str, Any], output_format: str = "pdf") -> bytes:
        """
        Generate resume by replacing placeholders in template
        
        Args:
            data: Dictionary containing all resume data
            output_format: "pdf" or "docx"
            
        Returns:
            bytes: Generated document as bytes
        """
        # Load the template
        doc = Document(self.template_path)
        
        # Define placeholder mappings
        replacements = self._prepare_replacements(data)
        
        # Replace placeholders in all paragraphs
        for paragraph in doc.paragraphs:
            self._replace_paragraph_text(paragraph, replacements)
        
        # Replace placeholders in all tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_paragraph_text(paragraph, replacements)
        
        # Save to temporary file
        temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_docx.name)
        temp_docx.close()
        
        try:
            if output_format.lower() == "pdf":
                # Convert to PDF using LibreOffice
                return self._convert_to_pdf(temp_docx.name)
            else:
                # Return DOCX bytes
                with open(temp_docx.name, 'rb') as f:
                    return f.read()
        finally:
            # Clean up temp file
            os.unlink(temp_docx.name)
    
    def _prepare_replacements(self, data: Dict[str, Any]) -> Dict[str, str]:
        """Prepare all text replacements based on data"""
        replacements = {}
        
        # Header information
        replacements['{{NAME}}'] = data.get('name', 'Your Name').upper()
        replacements['{{EMAIL}}'] = data.get('email', 'email@example.com')
        replacements['{{PHONE}}'] = data.get('phone', '(555) 123-4567')
        replacements['{{LINKEDIN}}'] = data.get('linkedin', 'linkedin.com/in/profile')
        replacements['{{GITHUB}}'] = data.get('github', 'github.com/username')
        
        # Summary
        replacements['{{SUMMARY}}'] = data.get('summary', 'Professional summary goes here...')
        
        # Skills sections
        skills = data.get('skills', {})
        replacements['{{SKILL_FULLSTACK}}'] = skills.get('fullstack', 'Node.js, TypeScript, React, Python')
        replacements['{{SKILL_CLOUD}}'] = skills.get('cloud', 'AWS, Docker, Kubernetes')
        replacements['{{SKILL_API}}'] = skills.get('api', 'RESTful APIs, GraphQL')
        replacements['{{SKILL_ARCHITECTURE}}'] = skills.get('architecture', 'Microservices, Serverless')
        replacements['{{SKILL_DATABASE}}'] = skills.get('database', 'MongoDB, PostgreSQL, Redis')
        replacements['{{SKILL_MONITORING}}'] = skills.get('monitoring', 'CloudWatch, Datadog')
        replacements['{{SKILL_TESTING}}'] = skills.get('testing', 'Jest, Cypress, CI/CD')
        replacements['{{SKILL_AI}}'] = skills.get('ai', 'AWS Bedrock, OpenAI APIs')
        replacements['{{CERTIFICATIONS}}'] = skills.get('certifications', 'AWS Certified')
        
        # Experience sections
        experiences = data.get('experience', [])
        for i in range(3):  # Support up to 3 experiences
            if i < len(experiences):
                exp = experiences[i]
                replacements[f'{{{{EXP{i+1}_TITLE}}}}'] = exp.get('title', 'Software Engineer')
                replacements[f'{{{{EXP{i+1}_COMPANY}}}}'] = exp.get('company', 'Company Name')
                replacements[f'{{{{EXP{i+1}_LOCATION}}}}'] = exp.get('location', 'City, State')
                replacements[f'{{{{EXP{i+1}_DURATION}}}}'] = exp.get('duration', '2020 - Present')
                
                # Format bullets
                bullets = exp.get('bullets', [])
                bullet_text = '\n'.join(f'● {bullet}' for bullet in bullets)
                replacements[f'{{{{EXP{i+1}_BULLETS}}}}'] = bullet_text
            else:
                # Remove unused experience sections
                replacements[f'{{{{EXP{i+1}_TITLE}}}}'] = ''
                replacements[f'{{{{EXP{i+1}_COMPANY}}}}'] = ''
                replacements[f'{{{{EXP{i+1}_LOCATION}}}}'] = ''
                replacements[f'{{{{EXP{i+1}_DURATION}}}}'] = ''
                replacements[f'{{{{EXP{i+1}_BULLETS}}}}'] = ''
        
        # Education
        education = data.get('education', {})
        replacements['{{EDU_DEGREE}}'] = education.get('degree', 'Bachelor of Science in Computer Science')
        replacements['{{EDU_SCHOOL}}'] = education.get('school', 'University Name')
        replacements['{{EDU_LOCATION}}'] = education.get('location', 'City, State')
        replacements['{{EDU_GRADUATION}}'] = education.get('graduation', 'May 2020')
        
        return replacements
    
    def _replace_paragraph_text(self, paragraph, replacements: Dict[str, str]):
        """Replace placeholders in a paragraph while preserving formatting"""
        if paragraph.text:
            for placeholder, replacement in replacements.items():
                if placeholder in paragraph.text:
                    # Replace in the paragraph text
                    paragraph.text = paragraph.text.replace(placeholder, replacement)
    
    def _convert_to_pdf(self, docx_path: str) -> bytes:
        """Convert DOCX to PDF using LibreOffice"""
        output_dir = tempfile.mkdtemp()
        
        try:
            # Use LibreOffice to convert to PDF
            cmd = [
                'soffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', output_dir,
                docx_path
            ]
            
            # Try different LibreOffice paths
            libreoffice_paths = [
                '/Applications/LibreOffice.app/Contents/MacOS/soffice',  # macOS
                'soffice',  # System PATH
                'libreoffice',  # Alternative name
            ]
            
            for lo_path in libreoffice_paths:
                try:
                    cmd[0] = lo_path
                    result = subprocess.run(cmd, capture_output=True, text=True)
                    if result.returncode == 0:
                        break
                except FileNotFoundError:
                    continue
            else:
                raise Exception("LibreOffice not found. Please install LibreOffice for PDF conversion.")
            
            # Find the generated PDF
            pdf_filename = os.path.splitext(os.path.basename(docx_path))[0] + '.pdf'
            pdf_path = os.path.join(output_dir, pdf_filename)
            
            if not os.path.exists(pdf_path):
                raise Exception(f"PDF conversion failed: {result.stderr}")
            
            # Read and return PDF bytes
            with open(pdf_path, 'rb') as f:
                return f.read()
                
        finally:
            # Clean up
            import shutil
            shutil.rmtree(output_dir, ignore_errors=True)


def format_tailored_resume_for_docx(tailored_resume: Dict[str, Any], 
                                  personal_info: Dict[str, Any],
                                  company: str,
                                  role: str) -> Dict[str, Any]:
    """
    Format the AI-generated tailored resume for DOCX template
    """
    # Start with personal info
    data = {
        'name': personal_info.get('name', os.environ.get('USER_NAME', 'Your Name')),
        'email': personal_info.get('email', os.environ.get('USER_EMAIL', 'email@example.com')),
        'phone': personal_info.get('phone', os.environ.get('USER_PHONE', '(555) 123-4567')),
        'linkedin': personal_info.get('linkedin', os.environ.get('USER_LINKEDIN', 'linkedin.com/in/profile')),
        'github': personal_info.get('github', os.environ.get('USER_GITHUB', 'github.com/username')),
        
        # Use tailored summary
        'summary': tailored_resume.get('summary', ''),
        
        # Format skills from tailored resume
        'skills': format_skills_section(tailored_resume.get('skills', {})),
        
        # Merge experience with tailored bullets
        'experience': merge_experience_data(
            personal_info.get('experience', []),
            tailored_resume.get('experience', [])
        ),
        
        # Use personal education info
        'education': personal_info.get('education', {})
    }
    
    return data


def format_skills_section(tailored_skills: Dict[str, List[str]]) -> Dict[str, str]:
    """Format skills into categories matching the template"""
    # Map tailored skills to template categories
    return {
        'fullstack': ', '.join(tailored_skills.get('languages', []) + tailored_skills.get('frameworks', [])),
        'cloud': ', '.join(tailored_skills.get('tools', [])),
        'api': ', '.join(tailored_skills.get('technical', [])),
        'architecture': 'Microservices, Serverless, Event-Driven',  # Can be enhanced
        'database': 'MongoDB, PostgreSQL, Redis, DynamoDB',  # Can be enhanced
        'monitoring': 'CloudWatch, Datadog, New Relic',  # Can be enhanced
        'testing': 'Jest, Cypress, CI/CD',  # Can be enhanced
        'ai': 'AWS Bedrock, OpenAI APIs, LangChain',  # Can be enhanced
        'certifications': 'AWS Certified Solutions Architect'  # From personal info
    }


def merge_experience_data(personal_experience: List[Dict], 
                         tailored_experience: List[Dict]) -> List[Dict]:
    """Merge personal experience info with tailored bullets"""
    merged = []
    
    for i, personal_exp in enumerate(personal_experience[:3]):  # Max 3 experiences
        exp_data = {
            'title': personal_exp.get('title', 'Software Engineer'),
            'company': personal_exp.get('company', 'Company Name'),
            'location': personal_exp.get('location', 'City, State'),
            'duration': personal_exp.get('duration', '2020 - Present'),
        }
        
        # Use tailored bullets if available
        if i < len(tailored_experience) and 'achievements' in tailored_experience[i]:
            exp_data['bullets'] = tailored_experience[i]['achievements']
        else:
            exp_data['bullets'] = personal_exp.get('bullets', [])
        
        merged.append(exp_data)
    
    return merged