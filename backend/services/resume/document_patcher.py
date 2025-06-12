"""
Document Patcher Service
Handles patching resume templates with generated content using placeholder replacement
"""

import os
import json
import logging
import re
from typing import Dict, List, Any, Optional
from docx import Document
from datetime import datetime

logger = logging.getLogger(__name__)

class DocumentPatcher:
    """Patch resume templates with generated content"""
    
    def __init__(self, data_dir: str = "data", output_dir: str = "output"):
        self.data_dir = data_dir
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        
    def patch_resume_template(
        self,
        sections: Dict[str, Any],
        session_data: Dict[str, Any],
        template_name: str = "placeholder_resume.docx"
    ) -> str:
        """
        Patch resume template with generated sections
        
        Args:
            sections: Generated section content
            session_data: Session context (company, role, etc.)
            template_name: Template file name
            
        Returns:
            Path to patched resume file
        """
        try:
            template_path = os.path.join(self.data_dir, template_name)
            if not os.path.exists(template_path):
                raise FileNotFoundError(f"Template not found: {template_path}")
            
            # Generate output filename
            company = session_data.get('company', 'Company')
            role = session_data.get('role', 'Position')
            timestamp = datetime.now().strftime("%Y%m%d_%H%M")
            
            safe_company = "".join(c for c in company if c.isalnum() or c in (' ', '-', '_')).strip()
            safe_role = "".join(c for c in role if c.isalnum() or c in (' ', '-', '_')).strip()
            
            filename = f"{safe_company}_{safe_role}_{timestamp}.docx".replace(" ", "_")
            output_path = os.path.join(self.output_dir, filename)
            
            # Load and preprocess document
            doc = Document(template_path)
            doc = self._preprocess_document(doc)
            
            # Apply aggressive 1-page optimization first
            from .enhancers.space_optimizer import SpaceOptimizer
            space_optimizer = SpaceOptimizer()
            optimized_sections = space_optimizer.optimize_for_one_page(sections)
            
            # Create placeholder mapping from optimized sections
            placeholder_mapping = self._create_placeholder_mapping(optimized_sections, session_data)
            
            # Extract placeholders from document
            placeholders = self._extract_placeholders(doc)
            logger.info(f"Found {len(placeholders)} placeholders in template")
            
            # Apply replacements
            replaced_count = 0
            for placeholder in placeholders:
                key = placeholder.strip("<>")
                
                if key in placeholder_mapping:
                    value = placeholder_mapping[key]
                    self._replace_placeholder(doc, placeholder, value)
                    replaced_count += 1
                    logger.debug(f"Replaced {placeholder} with content")
                else:
                    logger.warning(f"No content found for placeholder: {placeholder}")
            
            # Apply formatting
            self._apply_basic_formatting(doc)
            
            # Save patched document
            doc.save(output_path)
            
            logger.info(f"Patched resume saved: {output_path}")
            logger.info(f"Replaced {replaced_count}/{len(placeholders)} placeholders")
            
            return output_path
            
        except Exception as e:
            logger.error(f"Error patching resume template: {e}")
            raise
    
    def _create_placeholder_mapping(
        self, 
        sections: Dict[str, Any], 
        session_data: Dict[str, Any]
    ) -> Dict[str, str]:
        """Create mapping of placeholders to content"""
        
        mapping = {}
        
        # Personal info (static for now)
        mapping.update({
            "NAME": "Harsha Vippala",
            "EMAIL": "harsha.vippala@gmail.com",
            "PHONE": "(469) 509-1996",
            "LOCATION": "Irving, TX",
            "LINKEDIN": "linkedin.com/in/harshavippala"
        })
        
        # Generated sections
        if "summary" in sections:
            mapping["SUMMARY"] = sections["summary"]
            
        if "skills" in sections:
            mapping["SKILLS"] = sections["skills"]
            # Also map individual skill categories if available
            if " | " in sections["skills"]:
                skill_parts = sections["skills"].split(" | ")
                for i, part in enumerate(skill_parts):
                    if ":" in part:
                        category, skills = part.split(":", 1)
                        category_clean = category.strip()
                        skills_clean = skills.strip()
                        
                        # Map to specific template placeholders
                        if "Languages" in category_clean:
                            mapping["SKILLS_LANGUAGESFRAMEWORK"] = skills_clean
                        elif "Cloud" in category_clean:
                            mapping["SKILLS_CLOUDDEVOPS"] = skills_clean
                        elif "APIs" in category_clean:
                            mapping["SKILLS_APISINTEGRATION"] = skills_clean
                        elif "Architecture" in category_clean:
                            mapping["SKILLS_ARCHITECTUREDESIGN"] = skills_clean
                        elif "Databases" in category_clean:
                            mapping["SKILLS_DATABASESSTORAGE"] = skills_clean
                        elif "Monitoring" in category_clean:
                            mapping["SKILLS_MONITORINGOBSERVABILITY"] = skills_clean
                        elif "Testing" in category_clean:
                            mapping["SKILLS_TESTINGCICD"] = skills_clean
                        
                        # Also keep the generic mapping as fallback
                        category_key = category.strip().upper().replace(" ", "_").replace("&", "")
                        mapping[f"SKILLS_{category_key}"] = skills_clean
        
        # Experience bullets
        if "experience" in sections and isinstance(sections["experience"], list):
            bullets = sections["experience"]
            
            # Map to specific job placeholders (JOB1, JOB2, JOB3)
            for i, bullet in enumerate(bullets[:5]):  # Max 5 bullets for JOB1 (current position)
                mapping[f"JOB1_POINT{i+1}"] = bullet
            
            # Generic mapping for backward compatibility
            for i, bullet in enumerate(bullets[:6]):  # Max 6 bullets
                mapping[f"EXPERIENCE_BULLET_{i+1}"] = bullet
        
        # Add content for JOB2 and JOB3 from structured experience data
        from .resume_parser import ResumeParser
        resume_parser = ResumeParser()
        base_experiences = resume_parser.base_experiences
        
        if len(base_experiences) >= 2:
            # JOB2 - Liberty Mutual Senior Software Engineer
            job2_bullets = base_experiences[1].get('experience_highlights', [])
            for i, bullet in enumerate(job2_bullets[:5]):
                # Apply space optimization to fit format
                optimized_bullet = self._optimize_bullet_for_space(bullet)
                mapping[f"JOB2_POINT{i+1}"] = optimized_bullet
        
        if len(base_experiences) >= 3:
            # JOB3 - Liberty Mutual Software Engineer  
            job3_bullets = base_experiences[2].get('experience_highlights', [])
            for i, bullet in enumerate(job3_bullets[:3]):  # Only 3 bullets for JOB3
                # Apply space optimization to fit format
                optimized_bullet = self._optimize_bullet_for_space(bullet)
                mapping[f"JOB3_POINT{i+1}"] = optimized_bullet
        
        # Job context
        mapping["TARGET_COMPANY"] = session_data.get('company', '')
        mapping["TARGET_ROLE"] = session_data.get('role', '')
        
        return mapping
    
    def _extract_placeholders(self, doc: Document) -> List[str]:
        """Extract all placeholders from document"""
        
        placeholders = set()
        pattern = re.compile(r'<[A-Z0-9_&]+>')
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            for match in pattern.findall(para.text):
                placeholders.add(match)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for match in pattern.findall(para.text):
                            placeholders.add(match)
        
        return sorted(list(placeholders))
    
    def _preprocess_document(self, doc: Document) -> Document:
        """Preprocess document to handle split placeholders"""
        
        pattern = re.compile(r'<[A-Z0-9_&]+>')
        
        # Process paragraphs
        for para in doc.paragraphs:
            if any('<' in run.text for run in para.runs) and len(para.runs) > 1:
                text = para.text
                placeholders_in_para = pattern.findall(text)
                
                # If placeholders exist but are split across runs, merge them
                if placeholders_in_para and not any(ph in run.text for ph in placeholders_in_para for run in para.runs):
                    # Clear and recreate runs
                    for i in range(len(para.runs)):
                        para.runs[0]._element.getparent().remove(para.runs[0]._element)
                    para.add_run(text)
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if any('<' in run.text for run in para.runs) and len(para.runs) > 1:
                            text = para.text
                            placeholders_in_para = pattern.findall(text)
                            
                            if placeholders_in_para and not any(ph in run.text for ph in placeholders_in_para for run in para.runs):
                                for i in range(len(para.runs)):
                                    para.runs[0]._element.getparent().remove(para.runs[0]._element)
                                para.add_run(text)
        
        return doc
    
    def _replace_placeholder(self, doc: Document, placeholder: str, replacement: str):
        """Replace placeholder with content throughout document"""
        
        def process_paragraph(paragraph):
            """Process a single paragraph for placeholder replacement"""
            text = ''.join(run.text for run in paragraph.runs)
            if placeholder not in text:
                return False
            
            # Find runs containing the placeholder
            start_idx = end_idx = None
            joined = ''
            
            for i, run in enumerate(paragraph.runs):
                if start_idx is None and placeholder.startswith(run.text):
                    joined = run.text
                    start_idx = i
                    if joined == placeholder:
                        end_idx = i
                        break
                elif start_idx is not None:
                    joined += run.text
                    if joined == placeholder:
                        end_idx = i
                        break
            
            # Simple replacement if not split across runs
            if start_idx is None or end_idx is None:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, replacement)
                        return True
                return False
            
            # Handle split placeholder
            first_run = paragraph.runs[start_idx]
            before = ''.join(run.text for run in paragraph.runs[:start_idx])
            after = ''.join(run.text for run in paragraph.runs[end_idx+1:])
            new_text = before + replacement + after
            
            # Remove all runs and add new one
            for _ in range(len(paragraph.runs)):
                paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)
            
            new_run = paragraph.add_run(new_text)
            
            # Copy formatting from first run
            new_run.bold = first_run.bold
            new_run.italic = first_run.italic
            new_run.underline = first_run.underline
            if first_run.font.size:
                new_run.font.size = first_run.font.size
            if first_run.font.name:
                new_run.font.name = first_run.font.name
            
            return True
        
        # Process all paragraphs
        for para in doc.paragraphs:
            process_paragraph(para)
        
        # Process table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        process_paragraph(para)
    
    def _apply_basic_formatting(self, doc: Document):
        """Apply basic formatting to the document"""
        
        skill_indicators = ["Languages:", "Frameworks:", "Cloud:", "Databases:", "Architecture:"]
        
        for para in doc.paragraphs:
            if ':' in para.text and any(indicator in para.text for indicator in skill_indicators):
                # Bold formatting for skill headers
                for run in para.runs:
                    if ':' in run.text:
                        colon_index = run.text.find(':')
                        if colon_index >= 0:
                            before_colon = run.text[:colon_index+1]
                            after_colon = run.text[colon_index+1:]
                            
                            # Split into two runs: bold header and normal content
                            run.text = before_colon
                            run.bold = True
                            
                            if after_colon:
                                new_run = para.add_run(after_colon)
                                new_run.bold = False
                        break
    
    def create_diff_json(
        self, 
        sections: Dict[str, Any], 
        session_data: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Create diff JSON for tracking changes"""
        
        return {
            "metadata": {
                "company": session_data.get('company', ''),
                "role": session_data.get('role', ''),
                "generated_at": datetime.utcnow().isoformat(),
                "session_id": session_data.get('id', '')
            },
            "sections": sections,
            "placeholders": self._create_placeholder_mapping(sections, session_data)
        }
    
    def get_available_templates(self) -> List[str]:
        """Get list of available resume templates"""
        
        templates = []
        template_extensions = ['.docx', '.dotx']
        
        if os.path.exists(self.data_dir):
            for file in os.listdir(self.data_dir):
                if any(file.endswith(ext) for ext in template_extensions):
                    templates.append(file)
        
        return templates
    
    def _optimize_bullet_for_space(self, bullet: str) -> str:
        """Optimize bullet point for space constraints (1-page format)"""
        
        if len(bullet) <= 220:  # Target for 1-page format
            return bullet
        
        # Apply common abbreviations and optimizations
        optimized = bullet
        
        # Remove em-dashes first (user requirement)
        optimized = optimized.replace('—', '-').replace('–', '-')
        
        # Technical abbreviations (ATS-safe only)
        abbreviations = {
            'JavaScript': 'JS',
            'TypeScript': 'TS',
            'Kubernetes': 'K8s',
            # Removed risky abbreviations that ATS might not recognize:
            # 'development': 'dev',  # ATS might not recognize
            # 'implementation': 'impl',  # ATS might not recognize  
            # 'optimization': 'opt',  # ATS might not recognize
            # 'performance': 'perf',  # ATS might not recognize
            # 'architecture': 'arch',  # ATS might not recognize
        }
        
        for full, abbrev in abbreviations.items():
            optimized = optimized.replace(full, abbrev)
        
        # Remove redundant phrases (more aggressive for 1-page)
        redundant_phrases = [
            'for the ',
            'in order to ',
            'was able to ',
            'successfully ',
            ' that ',
            ' which ',
            ' enabling ',
            ' across ',
            ' utilizing ',
            ' demonstrating ',
            ' including ',
            ' approximately ',
            ' through '
        ]
        
        for phrase in redundant_phrases:
            optimized = optimized.replace(phrase, ' ')
        
        # Replace common long phrases with shorter equivalents
        phrase_replacements = [
            (' resulting in ', ' achieving '),
            (' contributing to ', ' boosting '),
            (' implementing ', ' adding '),
            (' over ', ' >'),
            (' and ', ' & '),
            (', ', ',')
        ]
        
        for old, new in phrase_replacements:
            optimized = optimized.replace(old, new)
        
        # Clean up extra spaces
        optimized = ' '.join(optimized.split())
        
        # If still too long, truncate at word boundary (stricter for 1-page)
        if len(optimized) > 220:
            words = optimized.split()
            while len(' '.join(words)) > 210 and words:  # Leave buffer
                words.pop()
            optimized = ' '.join(words)
        
        return optimized