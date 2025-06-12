"""
Template Preview Service
Handles preview generation and document template display
"""

import logging
import os
import json
from pathlib import Path
from typing import Dict, List, Any, Optional
from docx import Document
from .document_patcher import DocumentPatcher

logger = logging.getLogger(__name__)

class TemplatePreviewService:
    """Generate preview content from template documents"""
    
    def __init__(self):
        self.data_dir = Path(__file__).parent.parent.parent / "data"
        self.template_path = self.data_dir / "placeholder_resume.docx"
        self.base_resume_path = self.data_dir / "Harsha_Master.docx"
        self.document_patcher = DocumentPatcher()
        
        # Verify template files exist
        if not self.template_path.exists():
            logger.warning(f"Template not found: {self.template_path}")
        if not self.base_resume_path.exists():
            logger.warning(f"Base resume not found: {self.base_resume_path}")
    
    def get_template_preview(self) -> Dict[str, Any]:
        """
        Extract content from placeholder_resume.docx for preview display
        
        Returns:
            Dictionary with template structure and placeholders
        """
        try:
            if not self.template_path.exists():
                return self._get_fallback_template()
            
            doc = Document(self.template_path)
            template_content = {
                "sections": [],
                "placeholders": [],
                "structure": "template"
            }
            
            # Extract paragraphs and identify placeholders
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    # Check if this is a placeholder (contains <PLACEHOLDER>)
                    if '<' in text and '>' in text:
                        template_content["placeholders"].append(text)
                    
                    template_content["sections"].append({
                        "type": "paragraph",
                        "content": text,
                        "is_placeholder": '<' in text and '>' in text
                    })
            
            # Extract tables if any
            for table in doc.tables:
                table_content = []
                for row in table.rows:
                    row_content = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_content.append(cell_text)
                            if '<' in cell_text and '>' in cell_text:
                                template_content["placeholders"].append(cell_text)
                    if row_content:
                        table_content.append(row_content)
                
                if table_content:
                    template_content["sections"].append({
                        "type": "table",
                        "content": table_content,
                        "is_placeholder": False
                    })
            
            logger.info(f"Extracted template with {len(template_content['sections'])} sections")
            return template_content
            
        except Exception as e:
            logger.error(f"Error reading template: {e}")
            return self._get_fallback_template()
    
    def get_base_resume_content(self) -> Dict[str, Any]:
        """
        Extract content from Harsha_Master.docx as base resume
        
        Returns:
            Dictionary with structured resume content
        """
        try:
            if not self.base_resume_path.exists():
                return self._get_fallback_base_resume()
            
            doc = Document(self.base_resume_path)
            resume_content = {
                "sections": [],
                "structure": "base_resume"
            }
            
            current_section = None
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # Check if this looks like a section header
                if self._is_section_header(text):
                    if current_section:
                        resume_content["sections"].append(current_section)
                    
                    current_section = {
                        "type": "section",
                        "title": text,
                        "content": []
                    }
                else:
                    # Add to current section or create default
                    if not current_section:
                        current_section = {
                            "type": "section", 
                            "title": "Content",
                            "content": []
                        }
                    
                    current_section["content"].append(text)
            
            # Add final section
            if current_section:
                resume_content["sections"].append(current_section)
            
            logger.info(f"Extracted base resume with {len(resume_content['sections'])} sections")
            return resume_content
            
        except Exception as e:
            logger.error(f"Error reading base resume: {e}")
            return self._get_fallback_base_resume()
    
    def generate_live_preview(
        self, 
        sections: Dict[str, Any], 
        selected_keywords: List[str] = None
    ) -> Dict[str, Any]:
        """
        Generate a live preview with user's content
        
        Args:
            sections: Dictionary with generated resume sections
            selected_keywords: List of keywords to highlight
            
        Returns:
            Dictionary with preview content and highlighting
        """
        try:
            preview_content = {
                "sections": [],
                "highlighted_keywords": selected_keywords or [],
                "structure": "live_preview"
            }
            
            # Build preview sections
            if sections.get("summary"):
                preview_content["sections"].append({
                    "type": "summary",
                    "title": "Professional Summary",
                    "content": sections["summary"],
                    "highlighted": self._highlight_keywords(
                        sections["summary"], selected_keywords or []
                    )
                })
            
            if sections.get("skills"):
                preview_content["sections"].append({
                    "type": "skills",
                    "title": "Technical Skills", 
                    "content": sections["skills"],
                    "highlighted": self._highlight_keywords(
                        sections["skills"], selected_keywords or []
                    )
                })
            
            if sections.get("experience"):
                experience_content = sections["experience"]
                if isinstance(experience_content, list):
                    experience_text = "\n".join([f"• {bullet}" for bullet in experience_content])
                else:
                    experience_text = experience_content
                
                preview_content["sections"].append({
                    "type": "experience",
                    "title": "Work Experience",
                    "content": experience_text,
                    "highlighted": self._highlight_keywords(
                        experience_text, selected_keywords or []
                    )
                })
            
            return preview_content
            
        except Exception as e:
            logger.error(f"Error generating live preview: {e}")
            return {
                "sections": [],
                "highlighted_keywords": [],
                "structure": "error",
                "error": str(e)
            }
    
    def _is_section_header(self, text: str) -> bool:
        """Check if text looks like a section header"""
        section_indicators = [
            "SUMMARY", "PROFILE", "OBJECTIVE",
            "EXPERIENCE", "EMPLOYMENT", "WORK",
            "EDUCATION", "SKILLS", "TECHNICAL",
            "PROJECTS", "ACHIEVEMENTS", "CERTIFICATIONS"
        ]
        
        text_upper = text.upper()
        return (
            any(indicator in text_upper for indicator in section_indicators) or
            len(text.split()) <= 3 and text.isupper() or
            text.endswith(':')
        )
    
    def _highlight_keywords(self, text: str, keywords: List[str]) -> str:
        """Add highlighting markers for keywords in text"""
        if not keywords:
            return text
            
        highlighted = text
        for keyword in keywords:
            # Simple highlighting - replace with marked version
            highlighted = highlighted.replace(
                keyword, 
                f"<mark>{keyword}</mark>"
            )
        
        return highlighted
    
    def _get_fallback_template(self) -> Dict[str, Any]:
        """Fallback template structure when file can't be read"""
        return {
            "sections": [
                {
                    "type": "paragraph",
                    "content": "<PROFESSIONAL_SUMMARY>",
                    "is_placeholder": True
                },
                {
                    "type": "paragraph", 
                    "content": "Technical Skills: <TECHNICAL_SKILLS>",
                    "is_placeholder": True
                },
                {
                    "type": "paragraph",
                    "content": "Work Experience:",
                    "is_placeholder": False
                },
                {
                    "type": "paragraph",
                    "content": "<WORK_EXPERIENCE>",
                    "is_placeholder": True
                }
            ],
            "placeholders": [
                "<PROFESSIONAL_SUMMARY>",
                "<TECHNICAL_SKILLS>", 
                "<WORK_EXPERIENCE>"
            ],
            "structure": "fallback_template"
        }
    
    def _get_fallback_base_resume(self) -> Dict[str, Any]:
        """Fallback base resume when file can't be read"""
        return {
            "sections": [
                {
                    "type": "section",
                    "title": "Professional Summary",
                    "content": ["Software Engineer with 4+ years of experience"]
                },
                {
                    "type": "section",
                    "title": "Technical Skills",
                    "content": ["Python, JavaScript, Node.js, AWS, MongoDB"]
                },
                {
                    "type": "section",
                    "title": "Work Experience", 
                    "content": ["Software Engineer II at 7-Eleven", "Developed payment systems"]
                }
            ],
            "structure": "fallback_base"
        }