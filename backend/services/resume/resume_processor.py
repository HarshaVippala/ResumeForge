"""
Resume Processor Service
Handles resume preview generation and export functionality
"""

import os
import json
import logging
from typing import Dict, List, Any, Optional
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

class ResumeProcessor:
    """Process and export resume documents"""
    
    def __init__(self, output_dir: str = "output"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        
        # Base resume template structure
        self.base_structure = {
            "contact_info": {
                "name": "Your Name",
                "email": "your.email@domain.com", 
                "phone": "(555) 123-4567",
                "location": "City, State",
                "linkedin": "linkedin.com/in/yourprofile"
            },
            "sections": ["summary", "skills", "experience", "education", "projects"]
        }
    
    def create_preview(
        self,
        sections: Dict[str, Any],
        job_info: Dict[str, str]
    ) -> Dict[str, Any]:
        """
        Create preview data for frontend rendering
        
        Args:
            sections: Generated section content
            job_info: Job context (company, role)
            
        Returns:
            Preview data with formatted sections
        """
        try:
            preview = {
                "metadata": {
                    "company": job_info.get("company", ""),
                    "role": job_info.get("role", ""),
                    "generated_at": datetime.utcnow().isoformat(),
                    "total_sections": len(sections)
                },
                "formatted_sections": {},
                "stats": {
                    "total_words": 0,
                    "total_characters": 0,
                    "estimated_length": "1 page"
                }
            }
            
            total_words = 0
            total_chars = 0
            
            # Process each section
            for section_type, content in sections.items():
                formatted_section = self._format_section_for_preview(section_type, content)
                preview["formatted_sections"][section_type] = formatted_section
                
                # Calculate stats
                if isinstance(content, str):
                    words = len(content.split())
                    chars = len(content)
                elif isinstance(content, list):
                    words = sum(len(item.split()) for item in content)
                    chars = sum(len(item) for item in content)
                else:
                    words = chars = 0
                
                total_words += words
                total_chars += chars
            
            preview["stats"]["total_words"] = total_words
            preview["stats"]["total_characters"] = total_chars
            
            # Estimate page length
            if total_words > 500:
                preview["stats"]["estimated_length"] = "2 pages"
            elif total_words > 750:
                preview["stats"]["estimated_length"] = "2-3 pages"
            
            return preview
            
        except Exception as e:
            logger.error(f"Error creating preview: {e}")
            return {"error": str(e)}
    
    def _format_section_for_preview(self, section_type: str, content: Any) -> Dict[str, Any]:
        """Format section content for preview display"""
        
        if section_type == "summary":
            return {
                "type": "paragraph",
                "content": content,
                "style": "summary"
            }
        
        elif section_type == "skills":
            # Parse skills into categories if formatted
            if " | " in content:
                categories = content.split(" | ")
                return {
                    "type": "categorized_list",
                    "content": categories,
                    "style": "skills_categorized"
                }
            else:
                skills = [skill.strip() for skill in content.split(",")]
                return {
                    "type": "comma_list",
                    "content": skills,
                    "style": "skills_list"
                }
        
        elif section_type == "experience":
            if isinstance(content, list):
                return {
                    "type": "bullet_list",
                    "content": content,
                    "style": "experience_bullets"
                }
            else:
                return {
                    "type": "paragraph",
                    "content": content,
                    "style": "experience_text"
                }
        
        else:
            # Generic handling for other sections
            return {
                "type": "paragraph",
                "content": str(content),
                "style": "default"
            }
    
    def export_resume(
        self,
        sections: Dict[str, Any],
        format: str = "docx",
        template: str = "modern",
        job_info: Dict[str, str] = None
    ) -> str:
        """
        Export complete resume to specified format
        
        Args:
            sections: Generated section content
            format: Export format ('docx', 'pdf', 'latex')
            template: Template style
            job_info: Job context for filename
            
        Returns:
            File path of exported resume
        """
        try:
            if format.lower() == "docx":
                return self._export_docx(sections, template, job_info)
            elif format.lower() == "pdf":
                return self._export_pdf(sections, template, job_info)
            elif format.lower() == "latex":
                return self._export_latex(sections, template, job_info)
            else:
                raise ValueError(f"Unsupported export format: {format}")
                
        except Exception as e:
            logger.error(f"Error exporting resume: {e}")
            raise
    
    def _export_docx(
        self,
        sections: Dict[str, Any],
        template: str,
        job_info: Dict[str, str]
    ) -> str:
        """Export resume as DOCX document"""
        
        try:
            # Create document
            doc = Document()
            
            # Set margins
            for section in doc.sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.7)
                section.right_margin = Inches(0.7)
            
            # Add header with contact info
            self._add_docx_header(doc)
            
            # Add sections in order
            section_order = ["summary", "skills", "experience", "education", "projects"]
            
            for section_type in section_order:
                if section_type in sections:
                    self._add_docx_section(doc, section_type, sections[section_type])
            
            # Generate filename
            company = job_info.get("company", "Company") if job_info else "Resume"
            role = job_info.get("role", "Position") if job_info else ""
            timestamp = datetime.now().strftime("%Y%m%d_%H%M")
            
            safe_company = "".join(c for c in company if c.isalnum() or c in (' ', '-', '_')).strip()
            safe_role = "".join(c for c in role if c.isalnum() or c in (' ', '-', '_')).strip()
            
            filename = f"{safe_company}_{safe_role}_{timestamp}.docx".replace(" ", "_")
            file_path = os.path.join(self.output_dir, filename)
            
            # Save document
            doc.save(file_path)
            logger.info(f"Exported DOCX resume: {file_path}")
            
            return file_path
            
        except Exception as e:
            logger.error(f"Error exporting DOCX: {e}")
            raise
    
    def _add_docx_header(self, doc: Document):
        """Add header section with contact information"""
        
        # Name
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(self.base_structure["contact_info"]["name"])
        name_run.font.size = Inches(0.25)  # Large font for name
        name_run.bold = True
        
        # Contact info
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_info = self.base_structure["contact_info"]
        
        contact_line = f"{contact_info['email']} | {contact_info['phone']} | {contact_info['location']} | {contact_info['linkedin']}"
        contact_para.add_run(contact_line)
        
        # Add spacing
        doc.add_paragraph()
    
    def _add_docx_section(self, doc: Document, section_type: str, content: Any):
        """Add a section to the DOCX document"""
        
        # Section titles
        titles = {
            "summary": "PROFESSIONAL SUMMARY",
            "skills": "TECHNICAL SKILLS", 
            "experience": "PROFESSIONAL EXPERIENCE",
            "education": "EDUCATION",
            "projects": "PROJECTS"
        }
        
        # Add section title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(titles.get(section_type, section_type.upper()))
        title_run.bold = True
        title_run.font.size = Inches(0.15)
        
        # Add section content
        if section_type == "summary":
            para = doc.add_paragraph(content)
            para.space_after = Inches(0.1)
            
        elif section_type == "skills":
            para = doc.add_paragraph(content)
            para.space_after = Inches(0.1)
            
        elif section_type == "experience" and isinstance(content, list):
            # Add job title placeholder
            job_para = doc.add_paragraph()
            job_run = job_para.add_run("Software Engineer | Company Name | Date Range")
            job_run.bold = True
            
            # Add bullet points
            for bullet in content:
                bullet_para = doc.add_paragraph(bullet, style='ListBullet')
                bullet_para.space_after = Inches(0.05)
        
        # Add spacing after section
        doc.add_paragraph()
    
    def _export_pdf(self, sections: Dict[str, Any], template: str, job_info: Dict[str, str]) -> str:
        """Export resume as PDF (placeholder - requires additional libraries)"""
        
        # For now, create DOCX and suggest PDF conversion
        docx_path = self._export_docx(sections, template, job_info)
        
        # In a full implementation, you would convert DOCX to PDF here
        # using libraries like python-docx2pdf or reportlab
        
        logger.warning("PDF export not fully implemented - returning DOCX file")
        return docx_path
    
    def _export_latex(self, sections: Dict[str, Any], template: str, job_info: Dict[str, str]) -> str:
        """Export resume as LaTeX document"""
        
        try:
            # Generate LaTeX content
            latex_content = self._generate_latex_resume(sections, job_info)
            
            # Generate filename
            company = job_info.get("company", "Company") if job_info else "Resume"
            role = job_info.get("role", "Position") if job_info else ""
            timestamp = datetime.now().strftime("%Y%m%d_%H%M")
            
            safe_company = "".join(c for c in company if c.isalnum() or c in (' ', '-', '_')).strip()
            safe_role = "".join(c for c in role if c.isalnum() or c in (' ', '-', '_')).strip()
            
            filename = f"{safe_company}_{safe_role}_{timestamp}.tex".replace(" ", "_")
            file_path = os.path.join(self.output_dir, filename)
            
            # Write LaTeX file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(latex_content)
            
            logger.info(f"Exported LaTeX resume: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"Error exporting LaTeX: {e}")
            raise
    
    def _generate_latex_resume(self, sections: Dict[str, Any], job_info: Dict[str, str]) -> str:
        """Generate LaTeX resume content"""
        
        contact = self.base_structure["contact_info"]
        
        latex_template = f"""\\documentclass[11pt]{{article}}
\\usepackage[margin=0.7in]{{geometry}}
\\usepackage{{enumitem}}
\\usepackage{{hyperref}}
\\pagestyle{{empty}}

\\begin{{document}}

% Header
\\begin{{center}}
\\textbf{{\\Large {contact['name']}}} \\\\
{contact['email']} | {contact['phone']} | {contact['location']} | {contact['linkedin']}
\\end{{center}}

\\vspace{{0.2in}}

"""
        
        # Add sections
        if "summary" in sections:
            latex_template += f"""
\\section*{{Professional Summary}}
{self._escape_latex(sections['summary'])}

"""
        
        if "skills" in sections:
            latex_template += f"""
\\section*{{Technical Skills}}
{self._escape_latex(sections['skills'])}

"""
        
        if "experience" in sections and isinstance(sections['experience'], list):
            latex_template += """
\\section*{Professional Experience}
\\textbf{Software Engineer | Company Name | Date Range}
\\begin{itemize}[leftmargin=*]
"""
            for bullet in sections['experience']:
                latex_template += f"\\item {self._escape_latex(bullet)}\n"
            
            latex_template += "\\end{itemize}\n\n"
        
        latex_template += "\\end{document}"
        
        return latex_template
    
    def _escape_latex(self, text: str) -> str:
        """Escape special LaTeX characters"""
        
        escape_chars = {
            '&': '\\&',
            '%': '\\%', 
            '$': '\\$',
            '#': '\\#',
            '^': '\\textasciicircum{}',
            '_': '\\_',
            '{': '\\{',
            '}': '\\}',
            '~': '\\textasciitilde{}',
            '\\': '\\textbackslash{}'
        }
        
        for char, escaped in escape_chars.items():
            text = text.replace(char, escaped)
        
        return text
    
    def get_available_templates(self) -> List[str]:
        """Get list of available resume templates"""
        return ["modern", "classic", "minimal", "creative"]
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported export formats"""
        return ["docx", "latex", "pdf"]