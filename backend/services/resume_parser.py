"""
Resume Parser Service
Handles parsing and extraction of content from base resume documents
"""

import os
import json
import logging
from typing import Dict, List, Any, Optional
from docx import Document
import re

logger = logging.getLogger(__name__)

class ResumeParser:
    """Parse and extract content from resume documents"""
    
    def __init__(self, data_dir: str = "data"):
        self.data_dir = data_dir
        
        # Load base resume data
        self.base_profile = self._load_base_profile()
        self.base_experiences = self._load_base_experiences()
        
    def _load_base_profile(self) -> Dict[str, Any]:
        """Load base resume profile data"""
        try:
            profile_path = os.path.join(self.data_dir, "base_resume_profile.json")
            if os.path.exists(profile_path):
                with open(profile_path, 'r') as f:
                    return json.load(f)
        except Exception as e:
            logger.error(f"Error loading base profile: {e}")
        
        # Return default profile structure
        return {
            "programming_languages": [],
            "frameworks_libraries_tools": [],
            "databases": [],
            "cloud_devops_tools": [],
            "architecture_design_concepts": []
        }
    
    def _load_base_experiences(self) -> List[Dict[str, Any]]:
        """Load structured experience data"""
        try:
            exp_path = os.path.join(self.data_dir, "harsha_experiences_structured.json")
            if os.path.exists(exp_path):
                with open(exp_path, 'r') as f:
                    return json.load(f)
        except Exception as e:
            logger.error(f"Error loading base experiences: {e}")
        
        return []
    
    def get_base_resume_content(self) -> Dict[str, Any]:
        """Get structured base resume content"""
        
        # Generate skills section from profile
        skills_formatted = self._format_skills_section(self.base_profile)
        
        # Generate summary from experience
        summary = self._generate_base_summary()
        
        # Get latest experiences
        experience_bullets = self._get_experience_bullets()
        
        return {
            "personal_info": {
                "name": "Harsha Vippala",
                "email": "harsha.vippala@gmail.com",
                "phone": "(469) 509-1996",
                "location": "Irving, TX",
                "linkedin": "linkedin.com/in/harshavippala"
            },
            "summary": summary,
            "skills": skills_formatted,
            "experience": experience_bullets,
            "raw_profile": self.base_profile,
            "raw_experiences": self.base_experiences
        }
    
    def _format_skills_section(self, profile: Dict[str, Any]) -> str:
        """Format skills into readable section"""
        
        skills_parts = []
        
        # Map profile keys to display names
        skill_categories = {
            "programming_languages": "Languages",
            "frameworks_libraries_tools": "Frameworks & Tools", 
            "databases": "Databases",
            "cloud_devops_tools": "Cloud & DevOps",
            "architecture_design_concepts": "Architecture & Design"
        }
        
        for key, display_name in skill_categories.items():
            if key in profile and profile[key]:
                skills_list = profile[key]
                if skills_list:
                    skills_parts.append(f"{display_name}: {', '.join(skills_list)}")
        
        return " | ".join(skills_parts)
    
    def _generate_base_summary(self) -> str:
        """Generate base professional summary"""
        
        # Calculate total experience years
        current_year = 2024
        start_year = 2021  # From first job
        total_years = current_year - start_year
        
        # Get current role info
        current_job = self.base_experiences[0] if self.base_experiences else None
        
        if current_job:
            current_title = current_job.get("job_title", "Software Engineer")
            current_company = current_job.get("company_name", "")
        else:
            current_title = "Software Engineer"
            current_company = ""
        
        # Generate summary
        summary = f"Experienced {current_title} with {total_years}+ years of expertise in full-stack development, cloud architecture, and scalable system design. "
        summary += "Proven track record in building microservices, implementing payment integrations, and leading technical initiatives. "
        summary += "Skilled in Node.js, Python, AWS, and modern web technologies with strong focus on performance optimization and system reliability."
        
        return summary
    
    def _get_experience_bullets(self, limit: int = 6) -> List[str]:
        """Get top experience bullets from current role"""
        
        if not self.base_experiences:
            return []
        
        # Get bullets from most recent role
        current_job = self.base_experiences[0]
        bullets = current_job.get("experience_highlights", [])
        
        # Return top bullets
        return bullets[:limit]
    
    def extract_placeholders_from_template(self, template_path: str = None) -> List[str]:
        """Extract placeholders from resume template"""
        
        if not template_path:
            template_path = os.path.join(self.data_dir, "placeholder_resume.docx")
        
        if not os.path.exists(template_path):
            logger.error(f"Template not found: {template_path}")
            return []
        
        try:
            doc = Document(template_path)
            placeholders = set()
            pattern = re.compile(r'<[A-Z0-9_&]+>')
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                for match in pattern.findall(para.text):
                    placeholders.add(match)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for match in pattern.findall(para.text):
                                placeholders.add(match)
            
            return sorted(list(placeholders))
            
        except Exception as e:
            logger.error(f"Error extracting placeholders: {e}")
            return []
    
    def get_placeholder_mapping(self) -> Dict[str, str]:
        """Get mapping of placeholders to base content"""
        
        base_content = self.get_base_resume_content()
        
        mapping = {
            "NAME": base_content["personal_info"]["name"],
            "EMAIL": base_content["personal_info"]["email"],
            "PHONE": base_content["personal_info"]["phone"],
            "LOCATION": base_content["personal_info"]["location"],
            "LINKEDIN": base_content["personal_info"]["linkedin"],
            "SUMMARY": base_content["summary"],
            "SKILLS": base_content["skills"]
        }
        
        # Add experience placeholders
        if self.base_experiences:
            for i, job in enumerate(self.base_experiences[:3]):  # Top 3 jobs
                job_num = i + 1
                mapping[f"JOB{job_num}_TITLE"] = job.get("job_title", "")
                mapping[f"JOB{job_num}_COMPANY"] = job.get("company_name", "")
                mapping[f"JOB{job_num}_DATES"] = job.get("dates", "")
                mapping[f"JOB{job_num}_LOCATION"] = job.get("location", "")
                
                # Add experience bullets
                bullets = job.get("experience_highlights", [])
                for j, bullet in enumerate(bullets[:4]):  # Top 4 bullets per job
                    mapping[f"JOB{job_num}_BULLET{j+1}"] = bullet
        
        return mapping
    
    def get_skills_by_category(self) -> Dict[str, List[str]]:
        """Get skills organized by category"""
        return self.base_profile
    
    def search_experiences_by_keywords(self, keywords: List[str]) -> List[Dict[str, Any]]:
        """Find relevant experience bullets based on keywords"""
        
        relevant_experiences = []
        
        for job in self.base_experiences:
            job_match = {
                "company": job.get("company_name", ""),
                "title": job.get("job_title", ""),
                "dates": job.get("dates", ""),
                "matching_bullets": [],
                "technologies": job.get("technologies_used_in_role", [])
            }
            
            # Check each bullet for keyword matches
            for bullet in job.get("experience_highlights", []):
                bullet_lower = bullet.lower()
                matches = []
                
                for keyword in keywords:
                    if keyword.lower() in bullet_lower:
                        matches.append(keyword)
                
                if matches:
                    job_match["matching_bullets"].append({
                        "text": bullet,
                        "matched_keywords": matches,
                        "score": len(matches)
                    })
            
            # Only include jobs with matching bullets
            if job_match["matching_bullets"]:
                # Sort bullets by relevance score
                job_match["matching_bullets"].sort(key=lambda x: x["score"], reverse=True)
                relevant_experiences.append(job_match)
        
        return relevant_experiences
    
    def get_technologies_for_keywords(self, keywords: List[str]) -> List[str]:
        """Get relevant technologies based on keywords"""
        
        all_techs = set()
        
        # Add from profile categories
        for category_techs in self.base_profile.values():
            all_techs.update(category_techs)
        
        # Add from experience technologies
        for job in self.base_experiences:
            all_techs.update(job.get("technologies_used_in_role", []))
        
        # Filter technologies that match keywords
        relevant_techs = []
        for tech in all_techs:
            for keyword in keywords:
                if keyword.lower() in tech.lower() or tech.lower() in keyword.lower():
                    relevant_techs.append(tech)
                    break
        
        return list(set(relevant_techs))  # Remove duplicates